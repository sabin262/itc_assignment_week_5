from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


SAMPLE_DIR = Path("sample_leases")
SOURCE_GLOB = "*.txt"
PAGE_LINE_LIMIT = 58
PDF_LINE_WIDTH = 94


def main() -> None:
    for txt_path in sorted(SAMPLE_DIR.glob(SOURCE_GLOB)):
        text = txt_path.read_text(encoding="utf-8")
        docx_path = txt_path.with_suffix(".docx")
        pdf_path = txt_path.with_suffix(".pdf")

        write_docx(text, docx_path)
        write_pdf(text, pdf_path)
        print(f"{txt_path.name} -> {docx_path.name}, {pdf_path.name}")


def write_docx(text: str, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    normal_style.paragraph_format.space_after = Pt(6)

    pages = text.split("\f")
    for page_index, page_text in enumerate(pages):
        if page_index > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        for paragraph_text in split_paragraphs(page_text):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(paragraph_text)
            if is_title_or_page_heading(paragraph_text):
                run.bold = True
                paragraph.paragraph_format.space_before = Pt(8)
                paragraph.paragraph_format.space_after = Pt(8)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def split_paragraphs(text: str) -> list[str]:
    paragraphs: list[str] = []
    for block in text.replace("\r\n", "\n").split("\n\n"):
        cleaned = "\n".join(line.rstrip() for line in block.splitlines()).strip()
        if cleaned:
            paragraphs.append(cleaned)
    return paragraphs


def is_title_or_page_heading(text: str) -> bool:
    stripped = text.strip()
    if stripped.startswith("PAGE "):
        return True
    if len(stripped) <= 120 and stripped.upper() == stripped and any(ch.isalpha() for ch in stripped):
        return True
    return False


def write_pdf(text: str, output_path: Path) -> None:
    page_lines = paginate_pdf_lines(text)
    objects: list[bytes] = []

    page_object_numbers: list[int] = []
    next_object_number = 3
    for lines in page_lines:
        page_object_number = next_object_number
        content_object_number = next_object_number + 1
        page_object_numbers.append(page_object_number)
        next_object_number += 2

        content_stream = make_pdf_content_stream(lines)
        objects.append(
            (
                f"{page_object_number} 0 obj "
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 1 0 R >> >> "
                f"/Contents {content_object_number} 0 R >> endobj\n"
            ).encode("latin-1")
        )
        objects.append(
            (
                f"{content_object_number} 0 obj << /Length {len(content_stream)} >> stream\n"
            ).encode("latin-1")
            + content_stream
            + b"\nendstream endobj\n"
        )

    kids = " ".join(f"{number} 0 R" for number in page_object_numbers)
    header_objects = [
        b"1 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n",
        (
            f"2 0 obj << /Type /Pages /Kids [{kids}] /Count {len(page_object_numbers)} >> endobj\n"
        ).encode("latin-1"),
    ]
    catalog_object = (
        f"{next_object_number} 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
    ).encode("latin-1")

    all_objects = header_objects + objects + [catalog_object]
    write_pdf_objects(output_path, all_objects, root_object_number=next_object_number)


def paginate_pdf_lines(text: str) -> list[list[str]]:
    pages: list[list[str]] = []
    current: list[str] = []

    def flush_page() -> None:
        nonlocal current
        if current:
            pages.append(current)
            current = []

    for raw_page in text.split("\f"):
        for paragraph in split_paragraphs(raw_page):
            wrapped_lines = wrap(
                " ".join(paragraph.split()),
                width=PDF_LINE_WIDTH,
                replace_whitespace=False,
                drop_whitespace=True,
            ) or [""]
            for line in wrapped_lines:
                if len(current) >= PAGE_LINE_LIMIT:
                    flush_page()
                current.append(line)
            if len(current) >= PAGE_LINE_LIMIT:
                flush_page()
            current.append("")
        flush_page()

    if current:
        pages.append(current)
    return pages or [[""]]


def make_pdf_content_stream(lines: list[str]) -> bytes:
    commands = ["BT", "/F1 10 Tf", "50 760 Td", "13 TL"]
    for line in lines:
        commands.append(f"({escape_pdf_text(line)}) Tj")
        commands.append("T*")
    commands.append("ET")
    return "\n".join(commands).encode("latin-1", errors="replace")


def escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_pdf_objects(output_path: Path, objects: list[bytes], root_object_number: int) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_bytes in objects:
        offsets.append(len(pdf))
        pdf.extend(object_bytes)

    xref_at = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode("latin-1"))
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    pdf.extend(
        (
            f"trailer << /Size {len(objects) + 1} /Root {root_object_number} 0 R >>\n"
            f"startxref\n{xref_at}\n%%EOF\n"
        ).encode("latin-1")
    )
    output_path.write_bytes(pdf)


if __name__ == "__main__":
    main()
