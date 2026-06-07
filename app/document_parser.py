from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be converted to lease text."""


def extract_text_from_file(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentParseError(f"Unsupported file type '{extension}'. Use {supported}.")

    if extension == ".txt":
        text = _extract_txt(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    else:
        text = _extract_docx(content)

    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not cleaned:
        raise DocumentParseError(
            "No text could be extracted from the file. Scanned PDFs may require OCR."
        )
    return cleaned


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        page_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentParseError("Could not read the PDF file.") from exc
    return "\n\n".join(page_text)


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Could not read the DOCX file.") from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)
