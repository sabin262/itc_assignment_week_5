from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel, Field
from pypdf import PdfReader


app = FastAPI(
    title="Week 5 FastAPI Service",
    description="Backend service for Streamlit summarise, compare, and health pages.",
    version="1.0.0",
)


class SummaryRequest(BaseModel):
    text: str = Field(..., min_length=1, description="Text to summarise.")


class SummaryResponse(BaseModel):
    status: str
    summary: str
    original_length: int
    text: str
    filename: str | None = None


class CompareRequest(BaseModel):
    first_text: str = Field(..., min_length=1, description="First text value.")
    second_text: str = Field(..., min_length=1, description="Second text value.")


class CompareResponse(BaseModel):
    status: str
    are_equal: bool
    length_difference: int
    message: str
    first_text: str
    second_text: str
    first_filename: str | None = None
    second_filename: str | None = None


SUPPORTED_FILE_TYPES = {".txt", ".pdf", ".docx"}


def endpoint_status(name: str) -> dict[str, str]:
    return {"endpoint": name, "status": "available"}


def summarise_text(text: str, filename: str | None = None) -> SummaryResponse:
    cleaned = " ".join(text.split())
    words = cleaned.split()
    summary = " ".join(words[:30])

    if len(words) > 30:
        summary = f"{summary}..."

    return SummaryResponse(
        status="ok",
        summary=summary,
        original_length=len(text),
        text=text,
        filename=filename,
    )


def compare_text(
    first_text: str,
    second_text: str,
    first_filename: str | None = None,
    second_filename: str | None = None,
) -> CompareResponse:
    normalized_first = " ".join(first_text.lower().split())
    normalized_second = " ".join(second_text.lower().split())
    are_equal = normalized_first == normalized_second
    length_difference = abs(len(first_text) - len(second_text))

    return CompareResponse(
        status="ok",
        are_equal=are_equal,
        length_difference=length_difference,
        message="The submitted texts match." if are_equal else "The submitted texts differ.",
        first_text=first_text,
        second_text=second_text,
        first_filename=first_filename,
        second_filename=second_filename,
    )


def decode_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise HTTPException(status_code=400, detail="Could not decode the .txt file.")


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not extract text from the .pdf file.") from exc

    return "\n\n".join(page for page in pages if page)


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not extract text from the .docx file.") from exc

    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_FILE_TYPES:
        allowed = ", ".join(sorted(SUPPORTED_FILE_TYPES))
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Upload one of: {allowed}.")

    if extension == ".txt":
        text = decode_txt(file_bytes)
    elif extension == ".pdf":
        text = extract_pdf_text(file_bytes)
    else:
        text = extract_docx_text(file_bytes)

    if not text.strip():
        raise HTTPException(status_code=400, detail="No readable text was found in the uploaded file.")

    return text


async def get_uploaded_text(form: Any, field_name: str) -> tuple[str, str]:
    upload = form.get(field_name)

    if upload is None or not hasattr(upload, "filename") or not hasattr(upload, "read"):
        raise HTTPException(status_code=400, detail=f"Upload a file in the '{field_name}' field.")

    file_bytes = await upload.read()
    filename = upload.filename or "uploaded-file"
    text = extract_text_from_upload(filename, file_bytes)

    return filename, text


async def parse_summary_request(request: Request) -> tuple[str | None, str]:
    content_type = request.headers.get("content-type", "")

    if content_type.startswith("multipart/form-data"):
        form = await request.form()
        return await get_uploaded_text(form, "file")

    payload = SummaryRequest.model_validate(await request.json())
    return None, payload.text


async def parse_compare_request(request: Request) -> tuple[str | None, str, str | None, str]:
    content_type = request.headers.get("content-type", "")

    if content_type.startswith("multipart/form-data"):
        form = await request.form()
        first_filename, first_text = await get_uploaded_text(form, "first_file")
        second_filename, second_text = await get_uploaded_text(form, "second_file")
        return first_filename, first_text, second_filename, second_text

    payload = CompareRequest.model_validate(await request.json())
    return None, payload.first_text, None, payload.second_text


@app.get("/health")
def health() -> dict[str, object]:
    return {
        "service": "fastapi",
        "status": "ok",
        "endpoints": {
            "summarise": "available",
            "compare": "available",
            "health": "available",
        },
    }


@app.post("/summarise", response_model=SummaryResponse)
async def summarise(request: Request) -> SummaryResponse:
    filename, text = await parse_summary_request(request)
    return summarise_text(text=text, filename=filename)


@app.get("/summarise")
def summarise_status() -> dict[str, str]:
    return endpoint_status("/summarise")


@app.get("/compare")
def compare_status() -> dict[str, str]:
    return endpoint_status("/compare")


@app.post("/compare", response_model=CompareResponse)
async def compare(request: Request) -> CompareResponse:
    first_filename, first_text, second_filename, second_text = await parse_compare_request(request)
    return compare_text(
        first_text=first_text,
        second_text=second_text,
        first_filename=first_filename,
        second_filename=second_filename,
    )
